import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
import os
from matplotlib.ticker import MultipleLocator, FormatStrFormatter, FixedLocator
from datetime import datetime
from windrose import WindroseAxes

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL


class WeatherStation:
    """Класс для представления метеостанции и её метаданных"""

    def __init__(self, csv_path, city_name, case_city=None):
        """
        Инициализация метеостанции

        Args:
            csv_path (str): Путь к CSV файлу с данными
            city_name (str): Название города в именительном падеже
            case_city (str, optional): Название в предложном падеже ("в ...")
                                       Если не указано, используется city_name
        """
        self.csv_path = csv_path
        self.city_name = city_name
        self.case_city = case_city or city_name
        self.file_name = os.path.basename(csv_path)

    def get_metadata(self):
        """Получить метаданные для генерации отчетов"""
        return [
            f'wrose_met {self.city_name} за 10 лет.jpg',  # [0] для 10-летнего графика
            self.case_city,  # [1] город в предложном падеже
            f'{self.city_name} за период ',  # [2] префикс периода
            f'wrose {self.city_name}'  # [3] базовое имя для сохранения
        ]

    def __str__(self):
        return f"{self.city_name} ({self.file_name})"


def df_preparation(b_w, snow=True, metel=True):
    """Подготовка и фильтрация данных"""
    print("Подготовка данных...")

    # Создаем копию, чтобы избежать предупреждений
    b_w = b_w.copy()

    # Преобразование 'CALM' (штиль) в NaN
    b_w['wd'] = b_w['wd'].apply(lambda x: np.nan if x == 'CALM' else x)

    # Фильтрация по снегу
    if snow:
        # выбор только тех значений, где sn=2 (снег)
        b_w['sn'] = b_w['sn'].apply(lambda x: np.nan if x != 2 else x)

    # Фильтрация по скорости ветра
    if metel:
        b_w['ws'] = b_w['ws'].apply(lambda x: np.nan if x < 3 else x)

    # Удаление строк с NaN
    b_w = b_w.dropna()

    print(f"После фильтрации осталось {len(b_w)} записей")
    return b_w


def wind_rose(b_w, station, r_n, nw=False, total_count=None):
    """Построение розы ветров с правильными процентами из данных"""
    print("Построение розы ветров...")

    # Словарь для преобразования направлений ветра в градусы
    w_dir = {'N': 0, 'NNE': 22.5, 'NE': 45, 'ENE': 67.5, 'E': 90, 'ESE': 112.5,
             'SE': 135, 'SSE': 157.5, 'S': 180, 'SSW': 202.5, 'SW': 225,
             'WSW': 247.5, 'W': 270, 'WNW': 292.5, 'NW': 315, 'NNW': 337.5}

    # Преобразование буквенных обозначений в градусы
    b_w = b_w.copy()
    b_w['wd'] = b_w['wd'].apply(lambda x: w_dir[x] if x in w_dir else np.nan)
    b_w = b_w.dropna()

    if len(b_w) == 0:
        print("Нет данных для построения розы ветров после преобразования направлений")
        return

    # Создание оси для розы ветров
    fig = plt.figure(figsize=(10, 8))
    ax = WindroseAxes.from_ax(fig=fig)

    # Определяем бины для скорости ветра
    # Используем бины от 0 до максимальной скорости с шагом 1 м/с
    max_ws = int(b_w['ws'].max()) + 1
    bins = np.arange(0, max_ws + 1, 1)

    print(f"Диапазон скоростей: от {b_w['ws'].min():.1f} до {b_w['ws'].max():.1f} м/с")
    print(f"Бины: {bins}")

    # Построение графика в зависимости от типа
    if nw:
        # Только контур розы ветров
        ax.contour(b_w['wd'], b_w['ws'], bins=bins, colors='k', lw=2)
    else:
        if r_n == 0:
            # Черно-белая роза ветров
            ax.contourf(b_w['wd'], b_w['ws'], bins=bins, cmap=plt.cm.Greys)
            ax.contour(b_w['wd'], b_w['ws'], bins=bins, colors='k', lw=1)
        elif r_n == 1:
            # Цветная роза ветров с контуром
            ax.contourf(b_w['wd'], b_w['ws'], bins=bins, cmap=plt.cm.viridis)
            ax.contour(b_w['wd'], b_w['ws'], bins=bins, colors='k', lw=1)
        else:
            # Только контур (по умолчанию)
            ax.contour(b_w['wd'], b_w['ws'], bins=bins, colors='k', lw=2)

    # Настройка отображения
    ax.set_xticklabels(['В', 'СВ', 'С', 'СЗ', 'З', 'ЮЗ', 'Ю', 'ЮВ'],
                       fontsize=16, fontweight="bold")
    ax.set_theta_zero_location('E')

    # Добавляем легенду с процентами
    ax.set_legend(title='Скорость ветра, м/с', loc='center left', bbox_to_anchor=(1.1, 0.5))

    # ВЫЧИСЛЯЕМ РЕАЛЬНЫЕ ПРОЦЕНТЫ ИЗ ДАННЫХ ДЛЯ КРУГОВ

    # 1. Сначала получаем распределение данных по направлениям
    direction_counts = {}
    for direction in w_dir.values():
        # Считаем количество записей для каждого направления
        direction_data = b_w[b_w['wd'] == direction]
        direction_counts[direction] = len(direction_data)

    # 2. Находим максимальный процент среди всех направлений
    if total_count is None:
        total_count = len(b_w)

    max_percentage = 0
    for direction, count in direction_counts.items():
        percentage = (count / total_count * 100) if total_count > 0 else 0
        if percentage > max_percentage:
            max_percentage = percentage

    print(f"Максимальный процент среди направлений: {max_percentage:.1f}%")
    print(f"Общее количество данных: {total_count}")

    # 3. Настраиваем круги на основе реальных данных
    # Всегда используем 4 кольца
    circle_count = 4

    # Стандартные относительные позиции: 0.25, 0.5, 0.75, 1.0 (25%, 50%, 75%, 100%)
    standard_positions = np.array([0.25, 0.5, 0.75, 1.0])

    # Вычисляем фактические значения процентов на основе максимального
    if max_percentage > 0:
        # Масштабируем стандартные позиции под наши данные
        circle_values = standard_positions * max_percentage
    else:
        # Дефолтные значения, если нет данных
        circle_values = np.array([5, 10, 15, 20])

    # Округляем значения кругов
    circle_values = np.round(circle_values, 1)

    # Убеждаемся, что последнее кольцо соответствует максимальному проценту
    circle_values[-1] = max_percentage

    # Преобразуем проценты в позиции на графике
    ylim = ax.get_ylim()
    max_radius = ylim[1]

    if max_radius <= 0:
        max_radius = 1.0

    # Позиции кругов пропорциональны стандартным позициям
    circle_positions = standard_positions * max_radius

    # Устанавливаем круги
    ax.set_yticks(circle_positions)

    # Создаем метки для кругов
    circle_labels = []
    for value in circle_values:
        if value < 0.1:
            circle_labels.append(f'{value:.2f}%')
        elif value < 1:
            circle_labels.append(f'{value:.1f}%')
        else:
            if value.is_integer():
                circle_labels.append(f'{int(value)}%')
            else:
                circle_labels.append(f'{value:.1f}%')

    ax.set_yticklabels(circle_labels, fontsize=11, fontweight='bold')

    # Устанавливаем круги
    ax.set_yticks(circle_positions)

    # Создаем метки для кругов - проценты
    circle_labels = [f'{value:.1f}%' for value in circle_values]
    ax.set_yticklabels(circle_labels, fontsize=11, fontweight='bold')

    # # Добавляем процентные метки под кругами для ясности
    # ax.text(0.02, 0.98, 'Проценты указаны\nна кругах',
    #         transform=ax.transAxes,
    #         fontsize=10,
    #         verticalalignment='top',
    #         bbox=dict(boxstyle='round,pad=0.3', facecolor='white', alpha=0.8))

    # Сохранение графика
    metadata = station.get_metadata()
    plt.tight_layout()
    plt.savefig(metadata[3] + '.jpg', dpi=150, bbox_inches='tight')
    plt.close()

    print(f"Роза ветров сохранена как {metadata[3] + '.jpg'}")
    print(f"Круги отображают проценты: {', '.join(circle_labels)}")
    return


def obr_file(station, date_n, date_k, snow, metel, r_n, doc_name):
    """Основная функция обработки данных и генерации отчетов"""
    print(f"Обработка файла: {station.csv_path}")
    print(f"Период: с {date_n} по {date_k}")
    print(f"Условия: снег={snow}, ветер≥3м/с={metel}")

    # Временные метки для поиска ближайших измерений
    time_1 = [' 00:00', ' 03:00', ' 06:00', ' 09:00',
              ' 12:00', ' 15:00', ' 18:00', ' 21:00']

    # Загрузка CSV файла
    try:
        file_csv = pd.read_csv(station.csv_path, encoding='cp1251')
        print(f"Файл загружен, строк: {len(file_csv)}")
    except FileNotFoundError:
        print(f"Файл не найден: {station.csv_path}")
        return f"Ошибка: Файл {station.csv_path} не найден"
    except Exception as e:
        print(f"Ошибка при чтении файла: {e}")
        return f"Ошибка: Не удалось прочитать файл {station.csv_path}"

    # Проверяем необходимые колонки
    required_columns = ['dt_time', 'Wind_dir', 'wind_speed', 'precipitation']
    missing_columns = [col for col in required_columns if col not in file_csv.columns]
    if missing_columns:
        print(f"Отсутствуют колонки: {missing_columns}")
        return f"Ошибка: В файле отсутствуют колонки: {missing_columns}"

    # Получение метаданных станции
    metadata = station.get_metadata()

    # Поиск ближайшей доступной начальной даты
    original_date_n = date_n
    for i in range(8):
        date_1 = date_n + time_1[i]
        ind = file_csv.index[file_csv['dt_time'] == date_1].tolist()
        if len(ind) != 0:
            date_n = date_1
            print(f"Начальная дата найдена: {date_n}")
            break

    # Поиск ближайшей доступной конечной даты
    original_date_k = date_k
    for i in range(7, -1, -1):
        date_1 = date_k + time_1[i]
        ind = file_csv.index[file_csv['dt_time'] == date_1].tolist()
        if len(ind) != 0:
            date_k = date_1
            print(f"Конечная дата найдена: {date_k}")
            break

    # Находим индексы
    index_n = file_csv.loc[file_csv['dt_time'] == date_n].index
    index_k = file_csv.loc[file_csv['dt_time'] == date_k].index

    # Проверка корректности индексов
    if len(index_n) == 0 or len(index_k) == 0:
        print("Не удалось найти указанные даты в данных")
        return "Ошибка: Указанные даты не найдены в данных"

    # Определяем правильный порядок среза (от начала к концу)
    start_idx = min(index_n[0], index_k[0])
    end_idx = max(index_n[0], index_k[0])

    print(f"Индексы: start={start_idx}, end={end_idx}")

    # Выбор среза данных между указанными датами
    file_csv_daten_datek = file_csv.loc[start_idx:end_idx].copy()
    print(f"Выбран срез данных: {len(file_csv_daten_datek)} записей")

    # Создание DataFrame с нужными колонками
    b_wind = pd.DataFrame()
    b_wind['wd'] = file_csv_daten_datek['Wind_dir'].astype(str).str.strip()
    b_wind['ws'] = pd.to_numeric(file_csv_daten_datek['wind_speed'], errors='coerce')
    b_wind['sn'] = pd.to_numeric(file_csv_daten_datek['precipitation'], errors='coerce')

    print(f"Исходных данных: {len(b_wind)} записей")
    print(f"Направления ветра: {b_wind['wd'].unique()[:10]}...")
    print(f"Скорости ветра: от {b_wind['ws'].min():.1f} до {b_wind['ws'].max():.1f} м/с")

    # Подготовка данных (фильтрация)
    b_wind = df_preparation(b_wind, snow, metel)

    # Проверка, остались ли данные после фильтрации
    if b_wind.empty:
        print("Нет данных, соответствующих критериям фильтрации")
        return "Ошибка: Нет данных, соответствующих выбранным критериям"

    # Статистический анализ - считаем абсолютные значения
    res = b_wind.groupby(['ws', 'wd']).size().reset_index(name='count')
    total_count = res['count'].sum()
    print(f"Всего записей после фильтрации: {total_count}")

    # Создаем сводную таблицу с абсолютными значениями
    pivot_abs = pd.pivot_table(res, values='count', index=['ws'],
                               columns=['wd'], aggfunc='sum', fill_value=0)

    # ОКРУГЛЯЕМ АБСОЛЮТНЫЕ ЗНАЧЕНИЯ ДО 1 ЗНАКА ПОСЛЕ ЗАПЯТОЙ ПЕРЕД РАСЧЕТОМ ПРОЦЕНТОВ
    pivot_abs = pivot_abs.round(1)

    # Добавляем строку с итогами по направлениям
    direction_totals = pivot_abs.sum(axis=0)
    pivot_abs.loc['Всего'] = direction_totals

    # Добавляем столбец с итогами по скоростям (проценты)
    speed_totals = pivot_abs.sum(axis=1)
    pivot_abs['Всего,%'] = (speed_totals / total_count * 100).round(1)

    # Преобразуем абсолютные значения в проценты для отображения
    pivot_pct = (pivot_abs.iloc[:-1, :-1] / total_count * 100)

    # ИСПРАВЛЕНИЕ: Сначала округляем проценты, потом добавляем итоги
    pivot_pct = pivot_pct.round(1)

    # УБИРАЕМ МИКРО-ЗНАЧЕНИЯ МЕНЬШЕ 0.05 (делаем их 0.0)
    # Это решает проблему с 0.1 в строках, где все должно быть 0.0
    pivot_pct = pivot_pct.applymap(lambda x: 0.0 if abs(x) < 0.05 else x)

    # Добавляем строку с итогами по направлениям (в процентах)
    direction_pct = (direction_totals / total_count * 100).round(1)
    pivot_pct.loc['Всего,%'] = direction_pct

    # Добавляем столбец с итогами по скоростям
    # Пересчитываем, чтобы избежать погрешностей округления
    speed_pct_totals = pivot_pct.iloc[:-1, :].sum(axis=1)
    pivot_pct['Всего,%'] = speed_pct_totals.round(1)

    # Исправляем NaN в последней ячейке (пересечение "Всего,%" и "Всего,%")
    if 'Всего,%' in pivot_pct.index and 'Всего,%' in pivot_pct.columns:
        # Сумма всех процентов должна быть 100% (с учетом округления)
        total_sum = pivot_pct.iloc[:-1, :-1].sum().sum()
        pivot_pct.loc['Всего,%', 'Всего,%'] = total_sum

    print("\nСводная таблица (проценты):")
    print(pivot_pct)
    print(f"\nСумма всех процентов: {pivot_pct.iloc[:-1, :-1].sum().sum():.1f}%")

    # Исправленная строка - считаем сумму итоговой строки без последнего столбца
    if 'Всего,%' in pivot_pct.index:
        # Исключаем последний столбец 'Всего,%' из расчета
        direction_columns = [col for col in pivot_pct.columns if col != 'Всего,%']
        total_row_sum = pivot_pct.loc['Всего,%', direction_columns].sum()
        print(f"Сумма итоговой строки: {total_row_sum:.1f}%")

    # Построение розы ветров
    wind_rose(b_wind, station, r_n, False, total_count)

    # Формирование описания условий
    if snow:
        sn = 'Осадки в виде снега, '
    else:
        sn = 'Независимо от осадков,'

    if metel:
        sn = sn + 'ветер 3 и более м/с. '
    else:
        sn = sn + 'независимо от скорости ветра.'

    region = f"{metadata[1]}. Данные с {date_n} по {date_k}\n{sn}"

    # Генерация Word документа
    pivot_table_to_word(pivot_pct, pivot_abs, metadata, region, station.city_name, total_count, doc_name)


def pivot_table_to_word(pivot_pct, pivot_abs, metadata, region, city_name, total_count, doc_name):
    """Создание Word документа с таблицей и графиком"""
    print('Создание Word документа...')

    # Заголовки для направлений ветра
    headers = ['С', 'ССВ', 'СВ', 'СВС', 'В', 'ВЮВ', 'ЮВ', 'ЮЮВ', 'Ю',
               'ЮЮЗ', 'ЮЗ', 'ЗЮЗ', 'З', 'ЗСЗ', 'СЗ', 'ССЗ', 'Всего,%']

    # Обозначения направлений и их порядок
    directions_order = ['N', 'NNE', 'NE', 'ENE', 'E', 'ESE', 'SE', 'SSE', 'S',
                        'SSW', 'SW', 'WSW', 'W', 'WNW', 'NW', 'NNW']

    # Создаем словарь соответствия
    dir_mapping = {dir_name: headers[i] for i, dir_name in enumerate(directions_order)}

    # Создание документа
    document = Document()

    # Настройка стиля
    style = document.styles['Normal']
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(14)

    # Добавление заголовка
    p = document.add_paragraph('Роза ветров в ')
    p.alignment = 1  # Выравнивание по центру
    p.add_run(region)

    # Добавляем информацию об общем количестве данных
    p2 = document.add_paragraph(f'Всего обработано записей: {total_count}')
    p2.alignment = 1

    # Добавление изображения розы ветров
    image_path = metadata[3] + '.jpg'
    if os.path.exists(image_path):
        document.add_picture(image_path, width=Inches(6))
    else:
        document.add_paragraph(f"Изображение не найдено: {image_path}")

    # Создание новой страницы в альбомной ориентации для таблицы
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width

    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height

    # Уменьшение размера шрифта для таблицы
    style = document.styles['Normal']
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(9)

    # Создание таблицы
    table = document.add_table(rows=1, cols=18)
    table.style = 'Table Grid'

    # Заполнение заголовков таблицы
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Скорость ветра, м/с / Направление'

    for j in range(1, 18):
        hdr_cells[j].text = headers[j - 1]
        hdr_cells[j].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Получаем все направления из данных
    available_directions = [col for col in pivot_pct.columns if col in directions_order]
    print(f"Доступные направления в данных: {available_directions}")

    # Сортируем направления в правильном порядке
    available_directions.sort(key=lambda x: directions_order.index(x) if x in directions_order else 999)

    # Создаем словарь для быстрого доступа к данным по направлениям
    data_dict = {}
    for idx in pivot_pct.index:
        if idx != 'Всего,%':
            row_data = {}
            for dir_code in available_directions:
                if dir_code in pivot_pct.columns:
                    value = pivot_pct.loc[idx, dir_code]
                    # ИСПРАВЛЕНИЕ: Проверяем на микро-значения
                    if pd.isna(value) or abs(value) < 0.05:
                        row_data[dir_code] = 0.0
                    else:
                        row_data[dir_code] = float(value)
                else:
                    row_data[dir_code] = 0.0
            data_dict[idx] = row_data

    # Добавляем строку с итогами
    totals_row = {}
    for dir_code in available_directions:
        if dir_code in pivot_pct.columns and 'Всего,%' in pivot_pct.index:
            value = pivot_pct.loc['Всего,%', dir_code]
            if pd.isna(value) or abs(value) < 0.05:
                totals_row[dir_code] = 0.0
            else:
                totals_row[dir_code] = float(value)
        else:
            totals_row[dir_code] = 0.0
    data_dict['Всего,%'] = totals_row

    # Заполняем строки таблицы
    # Получаем все индексы кроме итоговой строки
    speed_indices = [idx for idx in pivot_pct.index if idx != 'Всего,%']

    # Сортируем скорости по возрастанию
    def sort_speed_key(x):
        try:
            return float(x)
        except:
            return float('inf')

    speed_indices.sort(key=sort_speed_key)

    # Добавляем итоговую строку в конец
    row_indices = speed_indices + ['Всего,%']

    for row_idx in row_indices:
        row_cells = table.add_row().cells

        # Заголовок строки
        if row_idx == 'Всего,%':
            row_cells[0].text = 'Всего,%'
        else:
            # Округляем скорость ветра до целого
            try:
                ws_value = int(float(row_idx))
                row_cells[0].text = str(ws_value)
            except:
                row_cells[0].text = str(row_idx)

        row_cells[0].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Заполняем данные по направлениям
        row_data = data_dict.get(row_idx, {})

        for j, dir_code in enumerate(directions_order):
            value = row_data.get(dir_code, 0.0)

            # Форматируем значение - ИСПРАВЛЕНИЕ: показываем 0.0 для значений < 0.05
            if abs(value) < 0.05:
                display_value = "0.0"
            else:
                display_value = f"{value:.1f}" if value != 0 else "0.0"

            row_cells[j + 1].text = display_value
            row_cells[j + 1].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.RIGHT

        # Заполняем последний столбец "Всего,%"
        if row_idx in pivot_pct.index and 'Всего,%' in pivot_pct.columns:
            total_value = pivot_pct.loc[row_idx, 'Всего,%']
            if pd.isna(total_value) or abs(total_value) < 0.05:
                row_cells[17].text = "0.0"
            else:
                row_cells[17].text = f"{float(total_value):.1f}"
        else:
            row_cells[17].text = "0.0"

        row_cells[17].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.RIGHT

    # Добавляем строку с проверкой суммы (только если есть итоговая строка)
    if 'Всего,%' in data_dict:
        # Считаем сумму по всем направлениям в итоговой строке
        direction_sum = sum([data_dict['Всего,%'].get(dir_code, 0.0) for dir_code in directions_order])

        check_row = table.add_row().cells
        check_row[0].text = "Проверка суммы:"
        check_row[0].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Объединяем ячейки для сообщения
        check_row[1].text = f"Сумма по направлениям: {direction_sum:.1f}%"
        for j in range(2, 17):
            check_row[1].merge(check_row[j])
        check_row[1].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Сохранение документа
    try:
        document.save(doc_name)
        print(f"Документ сохранен: {doc_name}")
    except Exception as e:
        print(f"Ошибка при сохранении документа: {e}")
