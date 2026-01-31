import os
from datetime import datetime

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt
from windrose import WindroseAxes

import weather_station


def make_report(
        station: weather_station.WeatherStation,
        date_from: datetime,
        date_to: datetime,
        snow_only: bool,
        wind_ge_3: bool,
        rose_type: int,
        doc_name: str,
) -> tuple[bool, str | None]:
    """
    Основная функция обработки данных и генерации отчетов.
    Возвращает: (успех, сообщение_об_ошибке_или_None)
    """
    try:
        # Валидация входных данных
        if date_from >= date_to:
            return False, f"Ошибка: дата 'с' ({date_from}) должна быть раньше даты 'по' ({date_to})"
        
        # Получение данных с фильтрацией
        try:
            data = station.get_data_for_rose_of_wind(
                date_from,
                date_to,
                snow_only,
                wind_ge_3,
            )
        except Exception as e:
            return False, f"Ошибка при загрузке данных из файла '{os.path.basename(station.csv_path)}': {str(e)}"
        
        # Проверка, остались ли данные после фильтрации
        if data.empty:
            conditions = []
            if snow_only:
                conditions.append("снег")
            if wind_ge_3:
                conditions.append("ветер ≥3 м/с")
            
            if conditions:
                cond_str = " + ".join(conditions)
                period = f"{date_from.strftime('%d.%m.%Y')} по {date_to.strftime('%d.%m.%Y')}"
                return False, f"Нет данных, удовлетворяющих условиям ({cond_str}) за период с {period}"
            else:
                period = f"{date_from.strftime('%d.%m.%Y')} по {date_to.strftime('%d.%m.%Y')}"
                return False, f"Нет данных за период с {period}"
        
        # 1. Абсолютная таблица
        try:
            abs_pivot = _create_absolute_pivot(data)
        except Exception as e:
            return False, f"Ошибка при создании таблицы распределения: {str(e)}"
        
        # 2. Процентная таблица
        try:
            total = abs_pivot.sum().sum()
            if total == 0:
                return False, "Ошибка: суммарное количество наблюдений равно нулю"
            
            rel_pivot = _create_percentage_pivot(abs_pivot, total_count=total)
        except Exception as e:
            return False, f"Ошибка при расчете процентов: {str(e)}"
        
        # Подготовка путей
        try:
            csv_dir = os.path.dirname(station.csv_path)
            metadata = station.get_metadata()
            image_full_path = os.path.join(csv_dir, f"{metadata[3]}.jpg")
        except Exception as e:
            return False, f"Ошибка при подготовке путей для сохранения: {str(e)}"
        
        # Построение розы ветров
        try:
            _draw_and_save_wind_rose_plot(
                data,
                rose_type,
                image_full_path,
                total_count=total,
            )
            
            # Проверка, что изображение сохранено
            if not os.path.exists(image_full_path):
                return False, f"Ошибка: файл изображения не был создан: {os.path.basename(image_full_path)}"
        except Exception as e:
            return False, f"Ошибка при построении розы ветров: {str(e)}"
        
        # Генерация Word документа
        try:
            _render_to_word(
                date_from,
                date_to,
                snow_only,
                wind_ge_3,
                metadata[1],
                rel_pivot,
                image_full_path,
                total,
                doc_name,
            )
            
            # Проверка, что документ сохранен
            if not os.path.exists(doc_name):
                return False, f"Ошибка: файл отчета не был сохранен: {os.path.basename(doc_name)}"
        except Exception as e:
            return False, f"Ошибка при создании Word-документа: {str(e)}"
        
        return True, None
    
    except Exception as e:
        return False, f"Неожиданная ошибка: {str(e)}"


def _draw_and_save_wind_rose_plot(data, rose_name, save_to, total_count=None):
    """
    Build Wind Rose Figure and save it to file.
    """
    # Словарь для преобразования направлений ветра в градусы
    w_dir = {
        "N": 0,
        "NNE": 22.5,
        "NE": 45,
        "ENE": 67.5,
        "E": 90,
        "ESE": 112.5,
        "SE": 135,
        "SSE": 157.5,
        "S": 180,
        "SSW": 202.5,
        "SW": 225,
        "WSW": 247.5,
        "W": 270,
        "WNW": 292.5,
        "NW": 315,
        "NNW": 337.5,
    }
    
    # Преобразование буквенных обозначений в градусы
    data = data.copy()
    data["wd"] = data["wd"].apply(lambda x: w_dir[x] if x in w_dir else np.nan)
    data = data.dropna()
    
    if len(data) == 0:
        raise ValueError("Нет данных для построения розы ветров после преобразования направлений")
    
    # Создание оси для розы ветров
    fig = plt.figure(figsize=(10, 8))
    ax = WindroseAxes.from_ax(fig=fig)
    
    # Определяем бины для скорости ветра
    max_ws = int(data["ws"].max())
    bins = np.arange(0, max_ws + 1, 1)
    
    # Построение графика в зависимости от типа
    if rose_name == 0:
        ax.contourf(data["wd"], data["ws"], bins=bins, cmap=plt.cm.Greys)
        ax.contour(data["wd"], data["ws"], bins=bins, colors="k", lw=1)
    elif rose_name == 1:
        ax.contourf(data["wd"], data["ws"], bins=bins, cmap=plt.cm.viridis)
        ax.contour(data["wd"], data["ws"], bins=bins, colors="k", lw=1)
    else:
        ax.contour(data["wd"], data["ws"], bins=bins, colors="k", lw=2)
    
    # Настройка отображения
    ax.set_xticklabels(
        ["В", "СВ", "С", "СЗ", "З", "ЮЗ", "Ю", "ЮВ"], fontsize=16, fontweight="bold"
    )
    ax.set_theta_zero_location("E")
    ax.set_legend(title="Скорость ветра, м/с", loc="center left", bbox_to_anchor=(1.1, 0.5))
    
    # Расчет и настройка кругов с процентами
    direction_counts = {}
    for direction in w_dir.values():
        direction_data = data[data["wd"] == direction]
        direction_counts[direction] = len(direction_data)
    
    if total_count is None or total_count <= 0:
        total_count = len(data)
        if total_count == 0:
            raise ValueError("Нет данных для расчета процентов")
    
    max_percentage = 0
    for count in direction_counts.values():
        percentage = (count / total_count * 100) if total_count > 0 else 0
        if percentage > max_percentage:
            max_percentage = percentage
    
    standard_positions = np.array([0.25, 0.5, 0.75, 1.0])
    if max_percentage > 0:
        circle_values = standard_positions * max_percentage
    else:
        circle_values = np.array([5, 10, 15, 20])
    
    circle_values = np.round(circle_values, 1)
    circle_values[-1] = max_percentage
    
    max_radius = ax.get_ylim()[1] or 1.0
    circle_positions = standard_positions * max_radius
    ax.set_yticks(circle_positions)
    
    circle_labels = [f"{value:.1f}%" for value in circle_values]
    ax.set_yticklabels(circle_labels, fontsize=11, fontweight="bold")
    
    # Сохранение графика
    try:
        plt.tight_layout()
        plt.savefig(save_to, dpi=150, bbox_inches="tight")
        plt.close()
    except Exception as e:
        plt.close()
        raise RuntimeError(f"Ошибка при сохранении изображения: {str(e)}")


def _create_absolute_pivot(data: pd.DataFrame):
    if data.empty:
        raise ValueError("Нет данных для создания таблицы")
    
    res = data.groupby(["ws", "wd"]).size().reset_index(name="count")
    pivot_abs = pd.pivot_table(
        res, values="count", index="ws", columns="wd", aggfunc="sum", fill_value=0
    )
    return pivot_abs


def _create_percentage_pivot(abs_pivot, total_count=None):
    if total_count is None:
        total_count = abs_pivot.sum().sum()
    
    if total_count <= 0:
        raise ValueError("Суммарное количество наблюдений равно нулю или отрицательно")
    
    pct_table = abs_pivot / total_count * 100
    direction_totals_pct = pct_table.sum(axis=0)
    pct_table.loc["Всего,%"] = direction_totals_pct
    
    speed_totals_pct = pct_table.iloc[:-1].sum(axis=1)
    pct_table["Всего,%"] = 0.0
    pct_table.loc[speed_totals_pct.index, "Всего,%"] = speed_totals_pct
    
    total_sum = pct_table.iloc[:-1, :-1].sum().sum()
    pct_table.loc["Всего,%", "Всего,%"] = total_sum
    
    return pct_table


# Константы направлений
DIRECTIONS_ORDER = [
    "N", "NNE", "NE", "ENE", "E", "ESE", "SE", "SSE",
    "S", "SSW", "SW", "WSW", "W", "WNW", "NW", "NNW",
]

DIRECTION_LABELS = [
    "С", "ССВ", "СВ", "СВС", "В", "ВЮВ", "ЮВ", "ЮЮВ",
    "Ю", "ЮЮЗ", "ЮЗ", "ЗЮЗ", "З", "ЗСЗ", "СЗ", "ССЗ",
]

assert len(DIRECTIONS_ORDER) == len(DIRECTION_LABELS), "Несоответствие количества направлений и меток"
DIR_MAPPING = dict(zip(DIRECTIONS_ORDER, DIRECTION_LABELS))
TABLE_HEADERS = ["Скорость ветра, м/с / Направление"] + DIRECTION_LABELS + ["Всего,%"]


def _format_value(value: float) -> str:
    """Форматирует значение для отображения: 0.0 если < 0.05, иначе с 1 знаком."""
    if pd.isna(value) or abs(value) < 0.05:
        return "0.0"
    return f"{value:.3f}"


def _get_sorted_speeds(index: pd.Index) -> list:
    """Сортирует индексы скоростей по числовому значению, нечисловые — в конец."""
    
    def key(x):
        try:
            return (0, float(x))
        except (ValueError, TypeError):
            return (1, str(x))
    
    return sorted([idx for idx in index if idx != "Всего,%"], key=key)


def _render_to_word(
        date_from,
        date_to,
        snow_only,
        wind_ge_3,
        case_city,
        rel_pivot: pd.DataFrame,
        rose_of_wind_img_path: str,
        total_count: int,
        doc_name: str,
):
    # Формирование описания условий (сохраняем оригинальную логику)
    if snow_only:
        sn = "Осадки в виде снега, "
    else:
        sn = "Независимо от осадков,"
    
    if wind_ge_3:
        sn = sn + "ветер 3 и более м/с. "
    else:
        sn = sn + "независимо от скорости ветра."
    
    region = f"{case_city}. Данные с {date_from} по {date_to}\n{sn}"
    
    # Создание документа
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    
    # Заголовок
    p = document.add_paragraph()
    p.alignment = WD_TABLE_ALIGNMENT.CENTER
    p.add_run(f"Роза ветров в {region}")
    
    # Общее количество записей
    p2 = document.add_paragraph()
    p2.alignment = WD_TABLE_ALIGNMENT.CENTER
    p2.add_run(f"Всего обработано записей: {total_count}")
    
    # Добавление изображения
    if os.path.exists(rose_of_wind_img_path):
        document.add_picture(rose_of_wind_img_path, width=Inches(6))
    else:
        document.add_paragraph(f"Изображение не найдено: {rose_of_wind_img_path}")
    
    # Альбомная ориентация для таблицы
    current_section = document.sections[-1]
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = current_section.page_height
    new_section.page_height = current_section.page_width
    style.font.size = Pt(9)
    
    # Создание таблицы
    n_cols = len(TABLE_HEADERS)
    table = document.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"
    
    # Заголовки
    for j, header in enumerate(TABLE_HEADERS):
        cell = table.rows[0].cells[j]
        cell.text = header
        cell.paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Подготовка данных
    available_dirs = [d for d in DIRECTIONS_ORDER if d in rel_pivot.columns]
    missing_dirs = set(DIRECTIONS_ORDER) - set(available_dirs)
    for d in missing_dirs:
        rel_pivot[d] = 0.0
    
    speed_rows = _get_sorted_speeds(rel_pivot.index)
    all_rows = speed_rows + (["Всего,%"] if "Всего,%" in rel_pivot.index else [])
    
    # Заполнение строк таблицы
    for row_label in all_rows:
        row_cells = table.add_row().cells
        
        if row_label == "Всего,%":
            row_cells[0].text = "Всего,%"
        else:
            try:
                ws_int = int(float(row_label))
                row_cells[0].text = str(ws_int)
            except (ValueError, TypeError):
                row_cells[0].text = str(row_label)
        row_cells[0].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for j, dir_code in enumerate(DIRECTIONS_ORDER):
            value = rel_pivot.loc[row_label, dir_code] if dir_code in rel_pivot.columns else 0.0
            row_cells[j + 1].text = _format_value(value)
            row_cells[j + 1].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.RIGHT
        
        total_col_val = rel_pivot.loc[row_label, "Всего,%"] if "Всего,%" in rel_pivot.columns else 0.0
        row_cells[-1].text = _format_value(total_col_val)
        row_cells[-1].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.RIGHT
    
    # Строка проверки суммы
    if "Всего,%" in rel_pivot.index:
        direction_sum = sum(
            rel_pivot.loc["Всего,%", d] for d in DIRECTIONS_ORDER if d in rel_pivot.columns
        )
        check_row = table.add_row().cells
        check_row[0].text = "Проверка суммы:"
        msg_cell = check_row[1]
        msg_cell.text = f"Сумма по направлениям: {direction_sum:.1f}%"
        for j in range(2, 17):
            msg_cell.merge(check_row[j])
        msg_cell.paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Сохранение документа
    try:
        document.save(doc_name)
    except Exception as e:
        raise RuntimeError(f"Ошибка при сохранении Word-документа: {str(e)}")